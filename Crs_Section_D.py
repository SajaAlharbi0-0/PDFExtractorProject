import json
import re
from docx import Document



# ==== UTILITY FUNCTIONS ====
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    return doc, full_text

# ==== SECTION D EXTRACTION ONLY ====
def extract_section_d(doc, text):
    data = {
        "Student Assessment": {
            "Assessment Activities": [],
            
        }
    }

    # ========== Assessment Activities Table ==========
    for table in doc.tables:
        if table.rows and len(table.rows) > 1:
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if len(headers) > 1 and "assessment activities" in headers[1]:
                for row in table.rows[1:]:
                    cells = row.cells
                    if len(cells) >= 3:
                        activity = cells[1].text.strip()
                        timing = cells[2].text.strip()
                        score = cells[3].text.strip() if len(cells) > 3 else ""
                        data["Student Assessment"]["Assessment Activities"].append({
                            "Activity": activity,
                            "Timing": timing,
                            "Score": score
                        })
                break  # Only the first matching table is needed

    

    return data


