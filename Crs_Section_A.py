from docx import Document
import re

def extract_section_a_and_info(docx_path):
    doc = Document(docx_path)
    full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

    data_structure = {
        "Course Info": {
            "Course Title": "",
            "Course Code": "",
            "Program": "",
            "Department": "",
            "Faculty": "",
            "University": "",
            "Version": "",
            "Last Revision Date": ""
        },
        "Sections": {
            "A": {
                "title": "General information about the course",
                "content": {
                    "1. Course Identification": {
                        "1. Credit hours": "",
                        "2. Course type": {
                            "A.": [],
                            "B.": []
                        },
                        "3. Level/year": "",
                        "4. Course general Description": "",
                        "5. Pre-requirements": "",
                        "6. Co-requirements": "",
                        "7. Course Main Objective(s)": ""
                    },
                    "2. Teaching mode": {
                        "Traditional classroom": {
                            "Contact Hours": None,
                            "Percentage": None
                        },
                        "E-learning": {
                            "Contact Hours": None,
                            "Percentage": None
                        },
                        "Hybrid • Traditional classroom • E-learning": {
                            "Contact Hours": None,
                            "Percentage": None
                        },
                        "Distance learning": {
                            "Contact Hours": None,
                            "Percentage": None
                        }
                    },
                    "3. Contact Hours": {
                        "Lectures": None,
                        "Laboratory/Studio": None,
                        "Field": None,
                        "Tutorial": None,
                        "Others (…)": None,
                        "Total": None
                    }
                }
            }
        }
    }

    # ✅ Extract Course Info values from tables
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip().replace("\n", " ").replace("\t", " ")
                value = row.cells[1].text.strip()
                if "Course Title" in key:
                    data_structure["Course Info"]["Course Title"] = value
                elif "Course Code" in key:
                    data_structure["Course Info"]["Course Code"] = value
                elif "Program" in key:
                    data_structure["Course Info"]["Program"] = value
                elif "Department" in key:
                    data_structure["Course Info"]["Department"] = value
                elif "Faculty" in key:
                    data_structure["Course Info"]["Faculty"] = value
                elif "University" in key:
                    data_structure["Course Info"]["University"] = value
                elif "Version" in key:
                    data_structure["Course Info"]["Version"] = value
                elif "Last Revision Date" in key:
                    data_structure["Course Info"]["Last Revision Date"] = value

    credit = re.search(r"1\. Credit hours:\s*\((.*?)\)", full_text)
    if credit:
        data_structure["Sections"]["A"]["content"]["1. Course Identification"]["1. Credit hours"] = credit.group(1).strip()

    if "☒ Required" in full_text:
        data_structure["Sections"]["A"]["content"]["1. Course Identification"]["2. Course type"]["B."].append("Required")
    if "☐ Elective" in full_text:
        data_structure["Sections"]["A"]["content"]["1. Course Identification"]["2. Course type"]["B."].append("Elective")
    if "☒ Faculty" in full_text:
        data_structure["Sections"]["A"]["content"]["1. Course Identification"]["2. Course type"]["A."].append("Faculty")

    level = re.search(r"3\. Level/year.*?:\s*(.+)", full_text)
    if level:
        data_structure["Sections"]["A"]["content"]["1. Course Identification"]["3. Level/year"] = level.group(1).strip()

    desc = re.search(r"4\. Course general Description:\s*(.+?)\s*5\.", full_text, re.DOTALL)
    if desc:
        data_structure["Sections"]["A"]["content"]["1. Course Identification"]["4. Course general Description"] = desc.group(1).strip()

    pre = re.search(r"5\. Pre-requirements.*?:\s*(.+?)\s*6\.", full_text, re.DOTALL)
    if pre:
        data_structure["Sections"]["A"]["content"]["1. Course Identification"]["5. Pre-requirements"] = pre.group(1).strip()

    co = re.search(r"6\. Co-requirements.*?:\s*(.+?)\s*7\.", full_text, re.DOTALL)
    if co:
        data_structure["Sections"]["A"]["content"]["1. Course Identification"]["6. Co-requirements"] = co.group(1).strip()

    obj = re.search(r"7\. Course Main Objective\(s\):\s*(.+?)\s*2\.", full_text, re.DOTALL)
    if obj:
        data_structure["Sections"]["A"]["content"]["1. Course Identification"]["7. Course Main Objective(s)"] = obj.group(1).strip()

    teaching_mode_match = re.search(r"1\s+Traditional classroom\s+(\d+)\s+(\d+%)", full_text)
    if teaching_mode_match:
        data_structure["Sections"]["A"]["content"]["2. Teaching mode"]["Traditional classroom"]["Contact Hours"] = int(teaching_mode_match.group(1))
        data_structure["Sections"]["A"]["content"]["2. Teaching mode"]["Traditional classroom"]["Percentage"] = teaching_mode_match.group(2)

    contact_hour_match = re.search(
        r"1\.\s*Lectures\s+(\d+)\s*2\.\s*Laboratory/Studio\s+(\d+)?\s*3\.\s*Field\s+(\d+)?\s*4\.\s*Tutorial\s+(\d+)?\s*5\.\s*Others.*?\s*(\d+)?\s*Total\s+(\d+)",
        full_text
    )
    if contact_hour_match:
        data_structure["Sections"]["A"]["content"]["3. Contact Hours"]["Lectures"] = int(contact_hour_match.group(1))
        data_structure["Sections"]["A"]["content"]["3. Contact Hours"]["Laboratory/Studio"] = int(contact_hour_match.group(2) or 0)
        data_structure["Sections"]["A"]["content"]["3. Contact Hours"]["Field"] = int(contact_hour_match.group(3) or 0)
        data_structure["Sections"]["A"]["content"]["3. Contact Hours"]["Tutorial"] = int(contact_hour_match.group(4) or 0)
        data_structure["Sections"]["A"]["content"]["3. Contact Hours"]["Others (…)"] = int(contact_hour_match.group(5) or 0)
        data_structure["Sections"]["A"]["content"]["3. Contact Hours"]["Total"] = int(contact_hour_match.group(6))

    return data_structure
