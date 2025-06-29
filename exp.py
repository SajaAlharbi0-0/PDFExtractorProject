import re
from docx import Document

def clean(text):
    return re.sub(r"\s+", " ", text.replace("\xa0", " ").replace("\n", " ")).strip()

def extract_to_json(file_path):
    doc = Document(file_path)
    full_text = "\n".join([clean(p.text) for p in doc.paragraphs])

    # === A. Field Experience Details ===
    field_experience_details = {
        "Credit Hours": "",
        "Level/Year": "",
        "Duration": {"Weeks": 0, "Days": 0, "Hours": 0},
        "Corequisite": "",
        "Delivery Mode": []
    }

    field_text = ""
    duration_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                field_text += cell_text + "\n"
                if any(k in cell_text.lower() for k in ["week", "day", "hour"]):
                    duration_text += cell_text + " "

    lines = field_text.splitlines()
    for i, line in enumerate(lines):
        line_lower = line.lower()
        if "credit hours" in line_lower:
            match = re.search(r"\(\s*(\d+)\s*credit", line_lower)
            if match:
                field_experience_details["Credit Hours"] = match.group(1)
        elif "level/year" in line_lower:
            match = re.search(r":\s*(.+)", line)
            if match:
                field_experience_details["Level/Year"] = match.group(1).strip()
        elif "corequisite" in line_lower:
            if i + 1 < len(lines):
                field_experience_details["Corequisite"] = lines[i + 1].strip()

        if re.search(r"[☒Xx]\s*in[- ]?person", line, re.I):
            field_experience_details["Delivery Mode"].append("In-person")
        if re.search(r"[☒Xx]\s*hybrid", line, re.I):
            field_experience_details["Delivery Mode"].append("Hybrid")
        if re.search(r"[☒Xx]\s*online", line, re.I):
            field_experience_details["Delivery Mode"].append("Online")

    weeks = re.search(r"\(?\s*(\d+)\s*\)?\s*weeks?", duration_text, re.I)
    days = re.search(r"\(?\s*(\d+)\s*\)?\s*days?", duration_text, re.I)
    hours = re.search(r"\(?\s*(\d+)\s*\)?\s*hours?", duration_text, re.I)
    if weeks: field_experience_details["Duration"]["Weeks"] = int(weeks.group(1))
    if days: field_experience_details["Duration"]["Days"] = int(days.group(1))
    if hours: field_experience_details["Duration"]["Hours"] = int(hours.group(1))

    # === B. CLOs ===
    def is_clo_table(table):
        if len(table.rows) == 0: return False
        header = [clean(cell.text).lower() for cell in table.rows[0].cells]
        keywords = ["code", "learning", "plo", "activities", "assessment", "responsibility"]
        return any(k in " ".join(header) for k in keywords)

    clos_data = {
        "1.0 Knowledge and understanding": [],
        "2.0 Skills": [],
        "3.0 Values, autonomy, and responsibility": []
    }

    for table in doc.tables:
        if is_clo_table(table):
            for row in table.rows[1:]:
                cells = row.cells
                if len(cells) < 6:
                    continue
                code = clean(cells[0].text)
                if not code or re.fullmatch(r"\d+\.0", code):
                    continue
                group_key = {
                    "1": "1.0 Knowledge and understanding",
                    "2": "2.0 Skills",
                    "3": "3.0 Values, autonomy, and responsibility"
                }.get(code[0], None)
                if group_key:
                    clos_data[group_key].append({
                        "Code": code,
                        "Course Learning Outcome": clean(cells[1].text),
                        "PLO Code": clean(cells[2].text),
                        "Teaching Strategies / Activities": clean(cells[3].text),
                        "Assessment Methods": clean(cells[4].text),
                        "Responsibility": clean(cells[5].text)
                    })

    # === C. Field Admin ===
    section_c_data = {
        "2. Distribution of Responsibilities for Field Experience Activities": [],
        "3. Field Experience Location Requirements": [],
        "5. Safety and Risk Management": []
    }

    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "activities" in headers[0] and "department" in headers[1]:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 6:
                    section_c_data["2. Distribution of Responsibilities for Field Experience Activities"].append({
                        "Activity": cells[0], "Department/College": cells[1], "Teaching Staff": cells[2],
                        "Student": cells[3], "Training Organization": cells[4], "Field Supervisor": cells[5]
                    })
        elif "suggested field experience locations" in headers[0]:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 3:
                    section_c_data["3. Field Experience Location Requirements"].append({
                        "Location": cells[0], "General Requirements": cells[1], "Special Requirements": cells[2]
                    })
        elif "potential risks" in headers[0]:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 3:
                    section_c_data["5. Safety and Risk Management"].append({
                        "Potential Risk": cells[0], "Safety Actions": cells[1], "Risk Management Procedures": cells[2]
                    })

    # === D. Flowchart ===
    def extract_flowchart_title_and_description(text: str) -> dict:
        pattern = re.compile(
            r"Field Experience Flowchart for Responsibility\s*(.*?)\n",
            re.IGNORECASE | re.DOTALL
        )
        match = re.search(pattern, text)
        description = match.group(1).strip() if match else ""
        return {
            "Title": "Field Experience Flowchart for Responsibility",
            "Description": description,
            "Note": "There is a flowchart image for this section."
        }

    # === E. Evaluation ===
    section_D_data = {"Training Quality Evaluation": []}
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "evaluation" in headers[0] and "evaluators" in headers[1]:
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 3:
                    section_D_data["Training Quality Evaluation"].append({
                        "Evaluation Area/Issue": cells[0],
                        "Evaluator": cells[1],
                        "Evaluation Method": cells[2]
                    })

    # === F. Approval ===
    def approval_data_from_tables(doc: Document) -> dict:
        for table in doc.tables:
            if len(table.rows) >= 3:
                try:
                    headers = [cell.text.strip().lower() for cell in table.columns[0].cells]
                    if "council" in headers[0] and "reference" in headers[1]:
                        return {
                            "Council/Committee": table.cell(0, 1).text.strip(),
                            "Reference No": table.cell(1, 1).text.strip(),
                            "Date": table.cell(2, 1).text.strip()
                        }
                except:
                    continue
        return {"Council/Committee": "", "Reference No": "", "Date": ""}

    # === General Info ===
    general_info_data = {}
    expected_labels = {
        "course title": "Course Title",
        "course code": "Course Code",
        "program": "Program",
        "department": "Department",
        "college": "College",
        "institution": "Institution",
        "field experience version number": "Version",
        "last revision date": "Last Revision Date"
    }

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 1:
                cell_text = row.cells[0].text.strip()
                if ":" in cell_text:
                    parts = cell_text.split(":", 1)
                    key_raw = parts[0].strip().lower()
                    value = parts[1].strip()
                    for key, label in expected_labels.items():
                        if key in key_raw:
                            general_info_data[label] = value

    return {
        "General Info": general_info_data,
        "A. Field Experience Details": field_experience_details,
        "B. CLOs": clos_data,
        "C. Field Experience Administration": section_c_data,
        "1. Field Experience Flowchart": extract_flowchart_title_and_description(full_text),
        "D. Evaluation": section_D_data,
        "E. Approval": approval_data_from_tables(doc),
    }
