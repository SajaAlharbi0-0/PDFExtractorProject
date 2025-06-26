import re
import json
import os
from docx import Document

def get_field_value(lines, field_name):
    pattern = rf"^{re.escape(field_name)}\s*:?\s*(.*)"
    for i, line in enumerate(lines):
        m = re.match(pattern, line)
        if m:
            val = m.group(1).strip()
            if val:
                return val
            if i+1 < len(lines):
                nxt = lines[i+1].strip()
                if nxt and not re.match(r"^\d+\.", nxt):
                    return nxt
    return ""

def extract_block(lines, start_num, stop_patterns):
    text, i, found = [], 0, False
    seen = set()  # track which lines we've already added

    start_re = rf"^{start_num}\."
    # find the start line
    while i < len(lines):
        if re.match(start_re, lines[i]):
            found = True
            parts = lines[i].split(":", 1)
            if len(parts) > 1 and parts[1].strip():
                part = parts[1].strip()
                if part not in seen:
                    text.append(part)
                    seen.add(part)
            i += 1
            break
        i += 1
    if not found:
        return ""

    # collect until any stop pattern, skipping duplicates
    while i < len(lines):
        if any(re.match(p, lines[i]) for p in stop_patterns):
            break
        part = lines[i].strip()
        if part and part not in seen:
            text.append(part)
            seen.add(part)
        i += 1

    return " ".join(text).strip()


def dedupe_text(text: str) -> str:
    parts = [p.strip() for p in text.split(":")]
    seen = set()
    deduped_parts = []
    for part in parts:
        lower = part.lower()
        if lower not in seen:
            deduped_parts.append(part)
            seen.add(lower)
    text = ": ".join(deduped_parts)

    if ":" in text:
        text = text.split(":", 1)[1].strip()

    text = re.sub(r'([^.]+\.)(?:\s*\1)+', r'\1', text)

    phrase_re = re.compile(r'\b((?:\w+\s+){0,4}\w+)\b(?:\s+\1\b)+', re.IGNORECASE)
    text = phrase_re.sub(r'\1', text)

    text = re.sub(r'\bNo\b$', '', text)
    return re.sub(r'\s{2,}', ' ', text).strip()

def common_prefix(s1, s2):
    result = ""
    for i in range(min(len(s1), len(s2))):
        if s1[i] != s2[i]:
            break
        result += s1[i]
    return result

def longest_repeated_substring(s):
    n = len(s)
    substrings = [s[i:] for i in range(n)]
    substrings.sort()
    lrs = ""
    for i in range(n - 1):
        temp = common_prefix(substrings[i], substrings[i + 1])
        if len(temp) > len(lrs):
            lrs = temp
    return lrs.strip()

def remove_lrs_repeats(text):
    lrs = longest_repeated_substring(text)
    if lrs and len(lrs) > 20:
        parts = text.split(lrs)
        cleaned = lrs + parts[-1] if len(parts) > 1 else text
        return cleaned.strip()
    return text

def extract_course_data(doc_path):
    doc = Document(doc_path)

    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    cells = [
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    ]
    raw_lines = [" ".join(line.strip().split()) for line in (paras + cells) if line.strip()]

    expanded = []
    for L in raw_lines:
        if "|" in L:
            for part in L.split("|"):
                part = part.strip()
                if part:
                    expanded.append(part)
        else:
            expanded.append(L)
    raw_lines = expanded

    print("=== DEBUG: Normalized lines ===")
    for idx, l in enumerate(raw_lines):
        print(f"{idx:03d}: {l}")
    print("=== END DEBUG ===\n")

    known = ["Course Title", "Course Code", "Program"]
    start = next((i for i, L in enumerate(raw_lines) if any(L.startswith(k) for k in known)), 0)
    lines = raw_lines[start:]

    data = {}

    single_fields = [
        "Course Title", "Course Code", "Program", "Department",
        "College", "Faculty", "Institution", "University", "Version", "Last Revision Date"
    ]
    for fld in single_fields:
        data[fld] = get_field_value(lines, fld)
    if data.get("College") and not data.get("Faculty"):
        data["Faculty"] = data["College"]
    if data.get("University") and not data.get("Institution"):
        data["Institution"] = data["University"]

    data["Credit Hours"] = ""
    for L in lines:
        if re.match(r"^1\.\s*Credit hours", L, re.I):
            data["Credit Hours"] = L.split(":", 1)[1].strip().strip("() ")
            break

    data["Course Type A"] = ""
    data["Course Type B"] = ""
    mode = None
    for L in lines:
        if re.match(r"^A\.", L): mode = "A"; continue
        if re.match(r"^B\.", L): mode = "B"; continue
        if mode in ("A", "B") and "☒" in L:
            data[f"Course Type {mode}"] = L.split("☒", 1)[1].strip()

    data["Level/Year"] = ""
    for L in lines:
        if re.match(r"^3\.\s*Level\/year", L, re.I):
            data["Level/Year"] = L.split(":", 1)[1].strip()
            break

    data["Course Description"] = extract_block(lines, 4, [r"^5\.", r"^Teaching"])
    data["Pre-requisites"] = extract_block(lines, 5, [r"^6\."])
    data["Co-requisites"] = extract_block(lines, 6, [r"^7\."])
    data["Course Main Objectives"] = extract_block(lines, 7, [r"^Teaching", r"^Mode"])

    for key in (
        "Course Description",
        "Pre-requisites",
        "Co-requisites",
        "Course Main Objectives",
    ):
        val = data.get(key, "")
        if isinstance(val, str) and val:
            val = dedupe_text(val)
            val = remove_lrs_repeats(val)
            data[key] = val

    data["Teaching Modes"], data["Activities"] = [], []

    for table in doc.tables:
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if "Mode of Instruction" in hdr:
            for row in table.rows[1:]:
                c = [cell.text.strip() for cell in row.cells]
                if len(c) >= 4:
                    data["Teaching Modes"].append({
                        "No": c[0],
                        "Mode": c[1],
                        "Contact Hours": c[2],
                        "Percentage": c[3]
                    })

    for table in doc.tables:
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if "Activity" not in hdr:
            continue
        for row in table.rows[1:]:
            c0 = row.cells[0].text.strip()
            c1 = row.cells[1].text.strip() if len(row.cells) > 1 else ""
            c2 = row.cells[2].text.strip() if len(row.cells) > 2 else ""
            data["Activities"].append({
                "No": c0,
                "Activity": c1,
                "Contact Hours": c2
            })
            if c0.lower() == "total":
                break
        break

    return data
